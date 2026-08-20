from __future__ import annotations

import csv
import io

from openpyxl import Workbook
from openpyxl.styles import Font
from docx import Document as DocxDocument

from constants import DEFAULT_ROLES
from models import ProjectDto


def _effective_hours(hours: float, buffer_percent: float) -> float:
    return hours * (1 + buffer_percent / 100)


def _fmt_date(iso: str | None) -> str:
    if not iso:
        return ""
    return iso[:10]


def _phase_breakdown_sheet(wb: Workbook, project: ProjectDto):
    sheet = wb.create_sheet("Phase Breakdown")
    headers = ["Phase", "Hours", "Buffer %", "Effective Hours", "Start", "End", "Rationale", "Dependencies"]
    sheet.append(headers)
    for cell in sheet[1]:
        cell.font = Font(bold=True)

    for phase in project.phases:
        sheet.append([
            phase.name,
            phase.hours,
            phase.bufferPercent,
            round(_effective_hours(phase.hours, phase.bufferPercent), 1),
            _fmt_date(phase.startDate),
            _fmt_date(phase.endDate),
            phase.rationale or "",
            phase.dependencies or "",
        ])

    total_row = ["Total", sum(p.hours for p in project.phases), "", round(sum(_effective_hours(p.hours, p.bufferPercent) for p in project.phases), 1), "", "", "", ""]
    sheet.append(total_row)
    for cell in sheet[sheet.max_row]:
        cell.font = Font(bold=True)

    widths = [26, 10, 10, 16, 14, 14, 50, 30]
    for i, w in enumerate(widths, start=1):
        sheet.column_dimensions[sheet.cell(row=1, column=i).column_letter].width = w


def _role_hours_sheet(wb: Workbook, project: ProjectDto):
    sheet = wb.create_sheet("Role Hours")
    roles = list(dict.fromkeys([*DEFAULT_ROLES, *project.customRoles, *[rh.role for p in project.phases for rh in p.roleHours]]))
    sheet.append(["Phase", *roles, "Total"])
    for cell in sheet[1]:
        cell.font = Font(bold=True)

    role_totals = {r: 0.0 for r in roles}
    for phase in project.phases:
        row = [phase.name]
        phase_total = 0.0
        for role in roles:
            hours = next((rh.hours for rh in phase.roleHours if rh.role == role), 0)
            row.append(hours)
            role_totals[role] += hours
            phase_total += hours
        row.append(phase_total)
        sheet.append(row)

    total_row = ["Total", *[role_totals[r] for r in roles], sum(role_totals.values())]
    sheet.append(total_row)
    for cell in sheet[sheet.max_row]:
        cell.font = Font(bold=True)

    sheet.column_dimensions["A"].width = 26


def _timeline_sheet(wb: Workbook, project: ProjectDto):
    sheet = wb.create_sheet("Timeline")
    sheet.append(["Phase", "Start", "End", "Duration (days)"])
    for cell in sheet[1]:
        cell.font = Font(bold=True)

    for phase in project.phases:
        duration = ""
        if phase.startDate and phase.endDate:
            from datetime import date as _date
            start = _date.fromisoformat(phase.startDate[:10])
            end = _date.fromisoformat(phase.endDate[:10])
            duration = (end - start).days + 1
        sheet.append([phase.name, _fmt_date(phase.startDate), _fmt_date(phase.endDate), duration])

    sheet.column_dimensions["A"].width = 26


def _stories_sheet(wb: Workbook, project: ProjectDto):
    sheet = wb.create_sheet("Stories")
    sheet.append(["Epic", "Story", "Description", "Acceptance Criteria", "Story Points", "Phase"])
    for cell in sheet[1]:
        cell.font = Font(bold=True)

    phase_name_by_id = {p.id: p.name for p in project.phases}
    for story in project.stories:
        sheet.append([
            story.epic,
            story.title,
            story.description,
            "\n".join(f"- {c}" for c in story.acceptanceCriteria),
            story.storyPoints,
            phase_name_by_id.get(story.phaseId, "") if story.phaseId else "",
        ])

    widths = [22, 40, 60, 60, 12, 22]
    for i, w in enumerate(widths, start=1):
        sheet.column_dimensions[sheet.cell(row=1, column=i).column_letter].width = w


def _tasks_sheet(wb: Workbook, project: ProjectDto):
    sheet = wb.create_sheet("Tasks")
    sheet.append(["Epic", "Story", "Task", "Hours", "Role"])
    for cell in sheet[1]:
        cell.font = Font(bold=True)

    for story in project.stories:
        for task in story.tasks:
            sheet.append([story.epic, story.title, task.title, task.hours, task.role])

    widths = [22, 40, 40, 10, 20]
    for i, w in enumerate(widths, start=1):
        sheet.column_dimensions[sheet.cell(row=1, column=i).column_letter].width = w


def _test_cases_sheet(wb: Workbook, project: ProjectDto):
    sheet = wb.create_sheet("Test Cases")
    sheet.append(["Test Case ID", "Story", "Title", "Precondition", "Steps", "Expected Result", "Priority", "Type"])
    for cell in sheet[1]:
        cell.font = Font(bold=True)

    counter = 0
    for story in project.stories:
        for tc in story.testCases:
            counter += 1
            sheet.append([
                f"TC-{counter:03d}",
                story.title,
                tc.title,
                tc.precondition or "",
                "\n".join(f"{i + 1}. {s}" for i, s in enumerate(tc.steps)),
                tc.expectedResult,
                tc.priority,
                tc.type,
            ])

    widths = [14, 40, 40, 30, 50, 40, 10, 14]
    for i, w in enumerate(widths, start=1):
        sheet.column_dimensions[sheet.cell(row=1, column=i).column_letter].width = w


def _workbook_bytes(wb: Workbook) -> bytes:
    wb.remove(wb["Sheet"])
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def build_estimate_workbook(project: ProjectDto) -> bytes:
    wb = Workbook()
    _phase_breakdown_sheet(wb, project)
    _role_hours_sheet(wb, project)
    _timeline_sheet(wb, project)
    return _workbook_bytes(wb)


def build_stories_workbook(project: ProjectDto) -> bytes:
    wb = Workbook()
    _stories_sheet(wb, project)
    _tasks_sheet(wb, project)
    return _workbook_bytes(wb)


def build_test_cases_workbook(project: ProjectDto) -> bytes:
    wb = Workbook()
    _test_cases_sheet(wb, project)
    return _workbook_bytes(wb)


def build_all_workbook(project: ProjectDto) -> bytes:
    wb = Workbook()
    _phase_breakdown_sheet(wb, project)
    _role_hours_sheet(wb, project)
    _timeline_sheet(wb, project)
    _stories_sheet(wb, project)
    _tasks_sheet(wb, project)
    _test_cases_sheet(wb, project)
    return _workbook_bytes(wb)


def build_estimate_csv(project: ProjectDto) -> str:
    buf = io.StringIO()
    writer = csv.writer(buf)
    writer.writerow(["Phase", "Hours", "Buffer %", "Effective Hours", "Start", "End", "Rationale", "Dependencies"])
    for phase in project.phases:
        writer.writerow([
            phase.name,
            phase.hours,
            phase.bufferPercent,
            round(_effective_hours(phase.hours, phase.bufferPercent), 1),
            _fmt_date(phase.startDate),
            _fmt_date(phase.endDate),
            phase.rationale or "",
            phase.dependencies or "",
        ])
    return buf.getvalue()


def build_stories_docx(project: ProjectDto) -> bytes:
    phase_name_by_id = {p.id: p.name for p in project.phases}
    epics: dict[str, list] = {}
    for story in project.stories:
        epics.setdefault(story.epic, []).append(story)

    doc = DocxDocument()
    doc.add_heading(project.name, level=0)
    doc.add_heading("Stories & Tasks", level=1)

    for epic_name, stories in epics.items():
        doc.add_heading(epic_name, level=2)
        for story in stories:
            doc.add_heading(f"{story.title} ({story.storyPoints} pts)", level=3)
            desc = doc.add_paragraph()
            desc.add_run(story.description).italic = True

            doc.add_paragraph("Acceptance Criteria:")
            for c in story.acceptanceCriteria:
                doc.add_paragraph(c, style="List Bullet")

            phase_name = phase_name_by_id.get(story.phaseId, "Unassigned") if story.phaseId else "Unassigned"
            doc.add_paragraph(f"Phase: {phase_name}")

            if story.tasks:
                table = doc.add_table(rows=1, cols=3)
                table.style = "Light Grid Accent 1"
                hdr = table.rows[0].cells
                hdr[0].text, hdr[1].text, hdr[2].text = "Task", "Hours", "Role"
                for task in story.tasks:
                    row = table.add_row().cells
                    row[0].text, row[1].text, row[2].text = task.title, str(task.hours), task.role
            doc.add_paragraph("")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
